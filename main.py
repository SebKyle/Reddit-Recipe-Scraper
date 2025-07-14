import praw
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

document = Document()
style = document.styles['Normal']
font = style.font

font.name = 'Times New Roman'
font.size = Pt(16)




reddit = praw.Reddit(
    client_id="Ye0hGHYVPuuMe_eunUGvXQ",
    client_secret="B3uzhCj5s2ZzPGozTx2KX8lyOIpoVw",
    user_agent="Recepie Script by u/HorseActual",
)

print(f"Connection Status: {reddit.read_only}")
for submission in reddit.subreddit("veganrecipes").hot(limit=30):
    if submission.stickied:
        continue
    RecipeAuthor= submission.author
    submission.comments.replace_more(limit=20)
    document.add_heading(submission.title , level=1)
    for comment in submission.comments.list():
        if comment.author == submission.author:
            document.add_paragraph(comment.body)
            break

    document.add_paragraph('')

document.save("Recipes.docx")